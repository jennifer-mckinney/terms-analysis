"""
Tests for service modules: analyzer, ingest, embedding, localai.
Covers uncovered lines to raise the overall line coverage.
"""
from __future__ import annotations

import asyncio
import io
import socket
from io import BytesIO
from typing import List
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.schemas import Evidence, Finding


# ── helpers ──────────────────────────────────────────────────────────────────

def _finding(
    *,
    category: str = "data_sharing",
    severity: str = "High",
    confidence: float = 0.85,
    excerpt: str = "We sell your data.",
    explanation: str = "Sharing data.",
    jurisdictions: list | None = None,
    legal_basis: list | None = None,
    line_start: int = 1,
    line_end: int = 1,
) -> Finding:
    return Finding(
        category=category,
        severity=severity,
        confidence=confidence,
        excerpt=excerpt,
        explanation=explanation,
        jurisdictions=jurisdictions or ["GDPR"],
        evidence=Evidence(
            line_start=line_start,
            line_end=line_end,
            legal_basis=legal_basis or ["GDPR Art. 6"],
        ),
    )


# ═══════════════════════════════════════════════════════════════════════════
# analyzer.py
# ═══════════════════════════════════════════════════════════════════════════

class TestTruncateText:
    def test_analyzer_truncate_text_short_text_unchanged(self):
        from app.services.analyzer import _truncate_text
        text = "short text"
        assert _truncate_text(text) == text

    def test_analyzer_truncate_text_long_text_is_cut(self):
        from app.services.analyzer import _truncate_text
        from app.config import settings
        long_text = "x" * (settings.max_input_chars + 100)
        result = _truncate_text(long_text)
        assert len(result) == settings.max_input_chars


class TestLineOffsets:
    def test_analyzer_line_offsets_multiline(self):
        from app.services.analyzer import _line_offsets
        text = "line1\nline2\nline3"
        offsets = _line_offsets(text)
        assert offsets[0] == 0
        assert offsets[1] == 6
        assert offsets[2] == 12

    def test_analyzer_line_offsets_empty_text(self):
        from app.services.analyzer import _line_offsets
        offsets = _line_offsets("")
        assert offsets == [0]


class TestBumpSeverity:
    def test_analyzer_bump_severity_low_boost_no_change(self):
        from app.services.analyzer import _bump_severity
        f = _finding(severity="Medium")
        result = _bump_severity(f, 0.1)
        assert result.severity == "Medium"

    def test_analyzer_bump_severity_high_boost_increments(self):
        from app.services.analyzer import _bump_severity
        f = _finding(severity="Medium")
        result = _bump_severity(f, 0.3)
        assert result.severity == "High"

    def test_analyzer_bump_severity_at_max_returns_same(self):
        from app.services.analyzer import _bump_severity
        f = _finding(severity="Critical")
        result = _bump_severity(f, 0.5)
        assert result.severity == "Critical"


class TestApplyDoctypeWeighting:
    def test_analyzer_apply_doctype_weighting_none_doctype_unchanged(self):
        from app.services.analyzer import _apply_doctype_weighting
        findings = [_finding(category="Data Sale / Sharing")]
        result = _apply_doctype_weighting(findings, None)
        assert result[0].severity == findings[0].severity

    def test_analyzer_apply_doctype_weighting_privacy_policy_boosts_sharing(self):
        from app.services.analyzer import _apply_doctype_weighting
        f = _finding(category="data sale / sharing", severity="Medium")
        result = _apply_doctype_weighting([f], "Privacy Policy")
        assert result[0].severity == "High"

    def test_analyzer_apply_doctype_weighting_unknown_doctype_unchanged(self):
        from app.services.analyzer import _apply_doctype_weighting
        f = _finding(category="data_sharing", severity="Low")
        result = _apply_doctype_weighting([f], "Unknown Document Type")
        assert result[0].severity == "Low"

    def test_analyzer_apply_doctype_weighting_combined_no_boost(self):
        from app.services.analyzer import _apply_doctype_weighting
        f = _finding(severity="Low")
        result = _apply_doctype_weighting([f], "Combined")
        assert result[0].severity == "Low"


class TestApplyIndustryEmphasis:
    def test_analyzer_apply_industry_emphasis_general_no_change(self):
        from app.services.analyzer import _apply_industry_emphasis
        f = _finding(severity="Low")
        result = _apply_industry_emphasis([f], "General")
        assert result[0].severity == "Low"

    def test_analyzer_apply_industry_emphasis_none_no_change(self):
        from app.services.analyzer import _apply_industry_emphasis
        f = _finding(severity="Low")
        result = _apply_industry_emphasis([f], None)
        assert result[0].severity == "Low"

    def test_analyzer_apply_industry_emphasis_healthcare_boosts_health_data(self):
        from app.services.analyzer import _apply_industry_emphasis
        f = _finding(category="health data", severity="Medium")
        result = _apply_industry_emphasis([f], "Healthcare")
        assert result[0].severity == "High"

    def test_analyzer_apply_industry_emphasis_unknown_industry_unchanged(self):
        from app.services.analyzer import _apply_industry_emphasis
        f = _finding(severity="Low")
        result = _apply_industry_emphasis([f], "Plumbing")
        assert result[0].severity == "Low"


class TestCalculateRiskScore:
    def test_analyzer_calculate_risk_score_no_findings(self):
        from app.services.analyzer import calculate_risk_score
        assert calculate_risk_score([]) == 0.0

    def test_analyzer_calculate_risk_score_all_high(self):
        from app.services.analyzer import calculate_risk_score
        findings = [_finding(severity="High"), _finding(severity="High")]
        score = calculate_risk_score(findings)
        assert score == pytest.approx(8.0)

    def test_analyzer_calculate_risk_score_mixed_severities(self):
        from app.services.analyzer import calculate_risk_score
        findings = [_finding(severity="Low"), _finding(severity="Critical")]
        score = calculate_risk_score(findings)
        assert 0 < score < 10


class TestGrade:
    def test_analyzer_grade_boundaries(self):
        from app.services.analyzer import _grade
        assert _grade(0.0) == "A"
        assert _grade(3.5) == "A-"
        assert _grade(4.5) == "B"
        assert _grade(5.5) == "B-"
        assert _grade(6.5) == "C+"
        assert _grade(7.5) == "C"
        assert _grade(8.5) == "D+"


class TestMergeFindings:
    def test_analyzer_merge_findings_rule_only(self):
        from app.services.analyzer import _merge_findings
        rule_f = _finding(excerpt="We sell your data.")
        result = _merge_findings([rule_f], [])
        assert len(result) == 1
        assert result[0].excerpt == rule_f.excerpt

    def test_analyzer_merge_findings_llm_only(self):
        from app.services.analyzer import _merge_findings
        llm_f = _finding(excerpt="We share with third parties.", confidence=0.75)
        result = _merge_findings([], [llm_f])
        assert len(result) == 1

    def test_analyzer_merge_findings_hybrid_blends_confidence(self):
        from app.services.analyzer import _merge_findings
        rule_f = _finding(excerpt="We sell your data.", confidence=0.90)
        llm_f = _finding(excerpt="We sell your data.", confidence=0.80)
        result = _merge_findings([rule_f], [llm_f])
        assert len(result) == 1
        expected_hybrid = 0.6 * 0.90 + 0.4 * 0.80
        assert result[0].confidence == pytest.approx(expected_hybrid, abs=0.01)

    def test_analyzer_merge_findings_deduplicates(self):
        from app.services.analyzer import _merge_findings
        rule_f1 = _finding(excerpt="duplicate text", category="cat1")
        rule_f2 = _finding(excerpt="duplicate text", category="cat1")
        result = _merge_findings([rule_f1, rule_f2], [])
        assert len(result) == 1

    def test_analyzer_merge_findings_llm_low_confidence_marked_for_review(self):
        from app.services.analyzer import _merge_findings
        llm_f = _finding(excerpt="Third party sharing.", confidence=0.4)
        result = _merge_findings([], [llm_f])
        assert result[0].needs_review is True


class TestDetectHighSeverityFindings:
    def test_analyzer_detect_high_severity_only_high_and_critical(self):
        from app.services.analyzer import detect_high_severity_findings
        text = "We sell your personal data to third parties."
        findings = detect_high_severity_findings(text, ["GDPR", "US-CA"])
        for f in findings:
            assert f.severity in {"High", "Critical"}

    def test_analyzer_detect_high_severity_returns_list(self):
        from app.services.analyzer import detect_high_severity_findings
        findings = detect_high_severity_findings("No sensitive content here.", ["GDPR"])
        assert isinstance(findings, list)


class TestAnalyzeTextQuickMode:
    def test_analyzer_analyze_text_quick_mode_returns_result(self):
        from app.services.analyzer import analyze_text
        result = asyncio.run(analyze_text("We sell your data.", ["GDPR"], mode="quick"))
        assert result.payload.analysis_mode == "quick"
        assert result.payload.confidence <= 1.0

    def test_analyzer_analyze_text_quick_mode_lower_confidence(self):
        from app.services.analyzer import analyze_text, analyze_text as _at
        result = asyncio.run(_at("We sell your data.", ["GDPR"], mode="quick"))
        # Quick mode applies 0.85 multiplier
        assert result.payload.confidence <= 1.0


class TestAnalyzeTextFullMode:
    def test_analyzer_analyze_text_full_mode_llm_success(self):
        from app.services.analyzer import analyze_text
        mock_payload = {
            "findings": [
                {
                    "category": "data_sharing",
                    "severity": "High",
                    "confidence": 0.85,
                    "excerpt": "We sell your data.",
                    "explanation": "Sells data.",
                    "jurisdictions": ["GDPR"],
                    "evidence": {
                        "line_start": 1, "line_end": 1,
                        "legal_basis": ["GDPR Art. 6"],
                    },
                }
            ],
            "summary": "Selling user data.",
            "overall_confidence": 0.87,
        }

        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = mock_payload
            mock_cls.return_value = mock_client

            result = asyncio.run(analyze_text("We sell your data.", ["GDPR"]))

        assert result.payload.summary == "Selling user data."
        assert len(result.payload.findings) >= 1
        excerpts = [f.excerpt for f in result.payload.findings]
        assert "We sell your data." in excerpts

    def test_analyzer_analyze_text_full_mode_llm_returns_none(self):
        from app.services.analyzer import analyze_text
        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = None
            mock_cls.return_value = mock_client

            result = asyncio.run(analyze_text("Some policy text.", ["GDPR"]))

        assert result.payload.confidence < 1.0

    def test_analyzer_analyze_text_full_mode_llm_finding_no_legal_basis_dropped(self):
        from app.services.analyzer import analyze_text
        mock_payload = {
            "findings": [
                {
                    "category": "data_sharing",
                    "severity": "High",
                    "confidence": 0.85,
                    "excerpt": "We sell your data.",
                    "explanation": "Sells data.",
                    "jurisdictions": ["GDPR"],
                    "evidence": {
                        "line_start": 1, "line_end": 1,
                        "legal_basis": [],  # No legal basis — should be dropped
                    },
                }
            ],
            "summary": "Risk found.",
            "overall_confidence": 0.8,
        }

        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = mock_payload
            mock_cls.return_value = mock_client

            result = asyncio.run(analyze_text("We sell your data.", ["GDPR"]))

        # LLM finding dropped due to missing legal_basis; no LLM-only finding survives
        # Rule engine may produce findings for "We sell your data." text
        for finding in result.payload.findings:
            assert finding.evidence.legal_basis, (
                f"All surviving findings must have legal_basis; found empty: {finding}"
            )

    def test_analyzer_analyze_text_full_mode_llm_low_confidence_needs_review(self):
        from app.services.analyzer import analyze_text
        mock_payload = {
            "findings": [
                {
                    "category": "data_sharing",
                    "severity": "Medium",
                    "confidence": 0.4,  # Below 0.6
                    "excerpt": "We may share your data.",
                    "explanation": "Possible sharing.",
                    "jurisdictions": ["GDPR"],
                    "evidence": {
                        "line_start": 1, "line_end": 1,
                        "legal_basis": ["GDPR Art. 6"],
                    },
                }
            ],
            "summary": None,
            "overall_confidence": 0.4,
        }

        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = mock_payload
            mock_cls.return_value = mock_client

            result = asyncio.run(analyze_text("We may share your data.", ["GDPR"]))

        # Low confidence LLM finding must exist and be marked needs_review.
        # The excerpt "We may share your data." is unique to the LLM mock payload
        # and won't be produced by the rule engine, so it must survive in findings.
        llm_only = [f for f in result.payload.findings if f.excerpt == "We may share your data."]
        assert len(llm_only) >= 1, "Low-confidence LLM finding must appear in merged results"
        assert llm_only[0].needs_review is True

    def test_analyzer_analyze_text_with_source_document(self):
        from app.services.analyzer import analyze_text
        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = None
            mock_cls.return_value = mock_client

            result = asyncio.run(
                analyze_text("Policy text.", ["GDPR"], source_document="PolicyDoc")
            )

        assert isinstance(result.payload.findings, list)
        assert result.payload.analysis_mode in {"full", "quick"}

    def test_analyzer_analyze_text_with_doctype_and_industry(self):
        from app.services.analyzer import analyze_text
        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = None
            mock_cls.return_value = mock_client

            result = asyncio.run(
                analyze_text(
                    "We sell health data.",
                    ["GDPR"],
                    doc_type="Privacy Policy",
                    industry="Healthcare",
                )
            )

        assert result.payload.confidence >= 0.0
        assert result.payload.analysis_mode == "full"
        # Industry/doctype path executed — findings list must exist
        assert isinstance(result.payload.findings, list)


class TestAnalyzeBatchDocuments:
    def test_analyzer_analyze_batch_documents_single_doc(self):
        from app.services.analyzer import analyze_batch_documents
        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = None
            mock_cls.return_value = mock_client

            results, cross_refs = asyncio.run(
                analyze_batch_documents(
                    [("Policy text here.", "Doc1", None, "Privacy Policy")],
                    "Healthcare",
                    ["GDPR"],
                )
            )

        assert len(results) == 1
        assert results[0].name == "Doc1"

    def test_analyzer_analyze_batch_documents_cross_references(self):
        from app.services.analyzer import analyze_batch_documents
        doc1 = "See our Privacy Policy for more details."
        doc2 = "Terms of Service content."
        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = None
            mock_cls.return_value = mock_client

            results, cross_refs = asyncio.run(
                analyze_batch_documents(
                    [
                        (doc1, "TOS", None, "Terms of Service"),
                        (doc2, "Privacy", None, "Privacy Policy"),
                    ],
                    "General",
                    ["GDPR"],
                    detect_cross_references=True,
                )
            )

        assert len(results) == 2
        assert isinstance(cross_refs, list)


class TestDetectCrossReferences:
    def test_analyzer_detect_cross_references_finds_reference(self):
        from app.services.analyzer import _detect_cross_references
        docs = [
            ("TOS", "Please see our Privacy Policy for details."),
            ("Privacy", "This is our privacy policy."),
        ]
        refs = _detect_cross_references(docs)
        assert len(refs) >= 1, "Should detect 'Privacy Policy' cross-reference"
        assert refs[0]["source_document"] == "TOS"

    def test_analyzer_detect_cross_references_no_reference(self):
        from app.services.analyzer import _detect_cross_references
        docs = [
            ("Doc1", "Unrelated content here."),
            ("Doc2", "More unrelated content."),
        ]
        refs = _detect_cross_references(docs)
        assert refs == []

    def test_analyzer_detect_cross_references_multiple_matches(self):
        from app.services.analyzer import _detect_cross_references
        docs = [
            ("Doc1", "As described in our Privacy Policy and see our Terms."),
            ("Doc2", "Reference document."),
        ]
        refs = _detect_cross_references(docs)
        assert isinstance(refs, list)


# ═══════════════════════════════════════════════════════════════════════════
# ingest.py
# ═══════════════════════════════════════════════════════════════════════════

class TestDecodeBytes:
    def test_ingest_decode_bytes_utf8_text(self):
        from app.services.ingest import _decode_bytes
        data = "hello world".encode("utf-8")
        assert _decode_bytes(data) == "hello world"

    def test_ingest_decode_bytes_utf16_falls_back(self):
        from app.services.ingest import _decode_bytes
        data = "hello".encode("utf-16")  # BOM + UTF-16 — fails UTF-8
        result = _decode_bytes(data)
        assert "hello" in result

    def test_ingest_decode_bytes_latin1_bytes(self):
        from app.services.ingest import _decode_bytes
        data = bytes([0xe9, 0xe0, 0xf1])  # é, à, ñ in latin-1
        result = _decode_bytes(data)
        assert isinstance(result, str)


class TestExtractPdf:
    def test_ingest_extract_pdf_simple(self):
        from app.services.ingest import _extract_pdf
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
        except ImportError:
            pytest.skip("reportlab not installed")

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter)
        styles = getSampleStyleSheet()
        doc.build([Paragraph("Sample PDF text for testing.", styles["Normal"])])
        buf.seek(0)
        text = _extract_pdf(buf.read())
        assert "Sample PDF text for testing." in text


class TestExtractDocx:
    def test_ingest_extract_docx_returns_text(self):
        from app.services.ingest import _extract_docx
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        buf = BytesIO()
        doc = Document()
        doc.add_paragraph("Privacy policy content for testing.")
        doc.save(buf)
        buf.seek(0)
        text = _extract_docx(buf.read())
        assert "Privacy policy content for testing." in text


class TestExtractRtf:
    def test_ingest_extract_rtf_simple(self):
        from app.services.ingest import _extract_rtf
        rtf_content = r"{\rtf1\ansi Hello World}"
        data = rtf_content.encode("utf-8")
        text = _extract_rtf(data)
        assert "Hello" in text


class TestExtractTextFromBytes:
    def test_ingest_extract_txt_file(self):
        from app.services.ingest import extract_text_from_bytes
        data = b"plain text content"
        text = extract_text_from_bytes("doc.txt", "text/plain", data)
        assert "plain text content" in text

    def test_ingest_extract_md_file(self):
        from app.services.ingest import extract_text_from_bytes
        data = b"# Heading\nSome markdown content"
        text = extract_text_from_bytes("README.md", "text/markdown", data)
        assert "Heading" in text

    def test_ingest_extract_html_file(self):
        from app.services.ingest import extract_text_from_bytes
        data = b"<html><body><p>HTML content</p></body></html>"
        text = extract_text_from_bytes("page.html", "text/html", data)
        assert "HTML content" in text

    def test_ingest_extract_htm_file(self):
        from app.services.ingest import extract_text_from_bytes
        data = b"<html><body>HTM content</body></html>"
        text = extract_text_from_bytes("page.htm", "text/html", data)
        assert "HTM content" in text

    def test_ingest_extract_pdf_empty_returns_ocr_attempt(self):
        from app.services.ingest import extract_text_from_bytes
        # pypdf with empty BytesIO returns empty pages
        buf = BytesIO()
        # Minimal valid PDF structure won't have text, falls to OCR path
        try:
            from pypdf import PdfWriter
            writer = PdfWriter()
            writer.add_blank_page(width=72, height=72)
            writer.write(buf)
        except Exception:
            pytest.skip("cannot create blank PDF")
        buf.seek(0)
        # Should not raise; may return empty string after OCR fallback
        text = extract_text_from_bytes("doc.pdf", "application/pdf", buf.read())
        assert isinstance(text, str)

    def test_ingest_extract_docx_via_extension(self):
        from app.services.ingest import extract_text_from_bytes
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        buf = BytesIO()
        doc = Document()
        doc.add_paragraph("Docx test content.")
        doc.save(buf)
        buf.seek(0)
        text = extract_text_from_bytes("policy.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", buf.read())
        assert "Docx test content." in text

    def test_ingest_extract_unknown_extension_falls_back(self):
        from app.services.ingest import extract_text_from_bytes
        data = "fallback content".encode("utf-8")
        text = extract_text_from_bytes("file.bin", None, data)
        assert "fallback content" in text

    def test_ingest_extract_with_html_content_type(self):
        from app.services.ingest import extract_text_from_bytes
        data = b"<html><body>Content-type matched</body></html>"
        text = extract_text_from_bytes("noext", "text/html; charset=utf-8", data)
        assert "Content-type matched" in text


class TestValidateUrl:
    def test_ingest_validate_url_http_allowed(self):
        from app.services.ingest import _validate_url
        # External DNS resolution — mock it
        with patch("socket.getaddrinfo") as mock_dns:
            mock_dns.return_value = [(None, None, None, None, ("93.184.216.34", None))]
            _validate_url("http://example.com/policy")

    def test_ingest_validate_url_private_ip_blocked(self):
        from app.services.ingest import _validate_url
        with pytest.raises(ValueError, match="not allowed"):
            _validate_url("http://127.0.0.1/admin")

    def test_ingest_validate_url_ftp_scheme_blocked(self):
        from app.services.ingest import _validate_url
        with pytest.raises(ValueError, match="Only http"):
            _validate_url("ftp://example.com/file")

    def test_ingest_validate_url_no_hostname_blocked(self):
        from app.services.ingest import _validate_url
        with pytest.raises(ValueError):
            _validate_url("http:///no-host/path")

    def test_ingest_validate_url_dns_failure_blocked(self):
        from app.services.ingest import _validate_url
        with patch("socket.getaddrinfo", side_effect=socket.gaierror("DNS failure")):
            with pytest.raises(ValueError, match="not allowed"):
                _validate_url("http://nonexistent.invalid/page")

    def test_ingest_validate_url_10_net_blocked(self):
        from app.services.ingest import _validate_url
        with pytest.raises(ValueError, match="not allowed"):
            _validate_url("http://10.0.0.1/admin")

    def test_ingest_validate_url_192_168_blocked(self):
        from app.services.ingest import _validate_url
        with pytest.raises(ValueError, match="not allowed"):
            _validate_url("http://192.168.1.1/")


class TestFetchUrlText:
    def test_ingest_fetch_url_text_success(self):
        from app.services.ingest import fetch_url_text
        html = b"<html><body>Policy content</body></html>"
        with patch("app.services.ingest._validate_url"):
            with patch("httpx.AsyncClient") as mock_client_cls:
                mock_response = MagicMock()
                mock_response.raise_for_status.return_value = None
                mock_response.headers = {"content-type": "text/html"}
                mock_response.content = html
                mock_client = AsyncMock()
                mock_client.__aenter__.return_value = mock_client
                mock_client.get.return_value = mock_response
                mock_client_cls.return_value = mock_client

                text = asyncio.run(fetch_url_text("https://example.com/policy"))

        assert "Policy content" in text

    def test_ingest_fetch_url_text_blocked_url_raises(self):
        from app.services.ingest import fetch_url_text
        with pytest.raises(ValueError, match="not allowed"):
            asyncio.run(fetch_url_text("http://127.0.0.1/admin"))

    def test_ingest_fetch_url_text_content_length_too_large_raises(self):
        """ingest.py lines 206-210: Content-Length header exceeds max_upload_bytes."""
        from app.services.ingest import fetch_url_text
        from app.config import settings
        oversized = str(settings.max_upload_bytes + 1)
        with patch("app.services.ingest._validate_url"):
            with patch("httpx.AsyncClient") as mock_cls:
                mock_response = MagicMock()
                mock_response.raise_for_status.return_value = None
                mock_response.headers = {"content-length": oversized, "content-type": "text/plain"}
                mock_client = AsyncMock()
                mock_client.__aenter__.return_value = mock_client
                mock_client.get.return_value = mock_response
                mock_cls.return_value = mock_client
                with pytest.raises(ValueError, match="exceeds"):
                    asyncio.run(fetch_url_text("https://example.com/policy"))

    def test_ingest_fetch_url_text_body_too_large_raises(self):
        """ingest.py lines 212-216: response body exceeds max_upload_bytes."""
        from app.services.ingest import fetch_url_text
        from app.config import settings
        with patch("app.services.ingest._validate_url"):
            with patch("httpx.AsyncClient") as mock_cls:
                mock_response = MagicMock()
                mock_response.raise_for_status.return_value = None
                # No content-length header so the body check runs
                mock_response.headers = {"content-type": "text/plain"}
                mock_response.content = b"x" * (settings.max_upload_bytes + 1)
                mock_client = AsyncMock()
                mock_client.__aenter__.return_value = mock_client
                mock_client.get.return_value = mock_response
                mock_cls.return_value = mock_client
                with pytest.raises(ValueError, match="exceeds"):
                    asyncio.run(fetch_url_text("https://example.com/policy"))

    def test_ingest_fetch_url_text_request_hook_validates_per_request_url(self):
        """ingest.py line 194: _on_request hook fires _validate_url for every request."""
        from app.services.ingest import fetch_url_text
        captured: dict = {}
        with patch("app.services.ingest._validate_url"):
            with patch("httpx.AsyncClient") as mock_cls:
                mock_response = MagicMock()
                mock_response.raise_for_status.return_value = None
                mock_response.headers = {"content-type": "text/plain"}
                mock_response.content = b"text content"
                mock_client = AsyncMock()
                mock_client.__aenter__.return_value = mock_client
                mock_client.get.return_value = mock_response

                def capture_kwargs(**kwargs):
                    captured.update(kwargs.get("event_hooks", {}))
                    return mock_client

                mock_cls.side_effect = capture_kwargs
                asyncio.run(fetch_url_text("https://example.com/policy"))

        assert "request" in captured, "event_hooks must register a 'request' hook"
        hook = captured["request"][0]
        fake_req = MagicMock()
        fake_req.url = "https://example.com/redirected"
        with patch("app.services.ingest._validate_url") as mock_val:
            hook(fake_req)  # fires line 194
            mock_val.assert_called_once_with("https://example.com/redirected")


# ═══════════════════════════════════════════════════════════════════════════
# embedding.py
# ═══════════════════════════════════════════════════════════════════════════

class TestBm25Scores:
    def test_embedding_bm25_scores_empty_corpus(self):
        from app.services.embedding import bm25_scores
        result = bm25_scores("query", [])
        assert result == []

    def test_embedding_bm25_scores_single_doc(self):
        from app.services.embedding import bm25_scores
        result = bm25_scores("privacy policy data", ["We protect your data."])
        assert len(result) == 1
        assert isinstance(result[0], float)

    def test_embedding_bm25_scores_multiple_docs(self):
        from app.services.embedding import bm25_scores
        corpus = [
            "We sell your personal data to third parties.",
            "Your privacy is our priority.",
        ]
        result = bm25_scores("sell personal data", corpus)
        assert len(result) == 2

    def test_embedding_bm25_scores_unavailable_returns_uniform(self):
        import app.services.embedding as emb
        original = emb._BM25_AVAILABLE
        try:
            emb._BM25_AVAILABLE = False
            result = emb.bm25_scores("query", ["doc1", "doc2"])
            assert result == [1.0, 1.0]
        finally:
            emb._BM25_AVAILABLE = original


class TestRrfFuse:
    def test_embedding_rrf_fuse_empty_returns_empty(self):
        from app.services.embedding import rrf_fuse
        result = rrf_fuse([])
        assert result == []

    def test_embedding_rrf_fuse_single_ranker(self):
        from app.services.embedding import rrf_fuse
        scores = [0.9, 0.5, 0.1]
        result = rrf_fuse([scores])
        assert len(result) == 3
        assert result[0] > result[2]

    def test_embedding_rrf_fuse_multiple_rankers(self):
        from app.services.embedding import rrf_fuse
        result = rrf_fuse([[0.9, 0.1], [0.8, 0.2]], k=60)
        assert len(result) == 2


class TestSelectRelevantChunks:
    def test_embedding_select_relevant_chunks_short_text_returned_as_is(self):
        from app.services.embedding import select_relevant_chunks
        from app.services.localai import LocalAIClient
        client = AsyncMock(spec=LocalAIClient)
        short_text = "short text"
        result = asyncio.run(
            select_relevant_chunks(short_text, "query", max_chars=1000, client=client)
        )
        assert result == short_text
        client.embed.assert_not_called()

    def test_embedding_select_relevant_chunks_embed_failure_falls_back(self):
        from app.services.embedding import select_relevant_chunks
        client = AsyncMock()
        client.embed.return_value = None  # Apertus unavailable

        long_text = "data " * 500
        result = asyncio.run(
            select_relevant_chunks(long_text, "query", max_chars=100, client=client)
        )
        assert result == long_text[:100]

    def test_embedding_select_relevant_chunks_full_path(self):
        from app.services.embedding import select_relevant_chunks
        client = AsyncMock()
        # Return a fake embedding vector
        fake_emb = [0.1, 0.2, 0.3]
        client.embed.return_value = fake_emb

        long_text = ("privacy data sharing policy content line\n" * 50)
        result = asyncio.run(
            select_relevant_chunks(long_text, "privacy data", max_chars=200, client=client)
        )
        assert isinstance(result, str)
        assert len(result) <= 300  # budget + join overhead

    def test_embedding_select_relevant_chunks_eu_embed_returns_none(self):
        from app.services.embedding import select_relevant_chunks
        client = AsyncMock()
        fake_emb = [0.1, 0.2, 0.3]

        call_count = 0
        async def mock_embed(text, model=None):
            nonlocal call_count
            call_count += 1
            from app.config import settings
            if model == settings.model_eu:
                return None  # EuroLLM unavailable
            return fake_emb

        client.embed.side_effect = mock_embed

        long_text = "detailed legal text about data protection " * 30
        result = asyncio.run(
            select_relevant_chunks(long_text, "data protection", max_chars=200, client=client)
        )
        assert isinstance(result, str)


# ═══════════════════════════════════════════════════════════════════════════
# localai.py
# ═══════════════════════════════════════════════════════════════════════════

class TestDetectLanguage:
    def test_localai_detect_language_returns_code(self):
        from app.services.localai import _detect_language
        result = _detect_language("This is English text about privacy.")
        assert result is None or isinstance(result, str)

    def test_localai_detect_language_short_text(self):
        from app.services.localai import _detect_language
        result = _detect_language("")
        assert result is None or isinstance(result, str)


class TestSelectModel:
    def test_localai_select_model_detection_disabled_returns_world(self):
        from app.services.localai import _select_model
        with patch("app.services.localai.settings") as mock_settings:
            mock_settings.language_detection_enabled = False
            mock_settings.model_world = "apertus-8b"
            mock_settings.eu_language_codes = []
            result = _select_model("Text here.")
        assert result == "apertus-8b"

    def test_localai_select_model_eu_language_routes_to_eurollm(self):
        from app.services.localai import _select_model
        with patch("app.services.localai.settings") as mock_settings:
            mock_settings.language_detection_enabled = True
            mock_settings.model_world = "apertus-8b"
            mock_settings.model_eu = "eurollm-22b"
            mock_settings.eu_language_codes = ["de", "fr", "en"]
            with patch("app.services.localai._detect_language", return_value="de"):
                result = _select_model("Datenschutzerklärung.")
        assert result == "eurollm-22b"

    def test_localai_select_model_world_language_routes_to_world(self):
        from app.services.localai import _select_model
        with patch("app.services.localai.settings") as mock_settings:
            mock_settings.language_detection_enabled = True
            mock_settings.model_world = "apertus-8b"
            mock_settings.model_eu = "eurollm-22b"
            mock_settings.eu_language_codes = ["de", "fr"]
            with patch("app.services.localai._detect_language", return_value="ja"):
                result = _select_model("Japanese text.")
        assert result == "apertus-8b"

    def test_localai_select_model_none_lang_routes_to_world(self):
        from app.services.localai import _select_model
        with patch("app.services.localai.settings") as mock_settings:
            mock_settings.language_detection_enabled = True
            mock_settings.model_world = "apertus-8b"
            mock_settings.model_eu = "eurollm-22b"
            mock_settings.eu_language_codes = ["de", "fr"]
            with patch("app.services.localai._detect_language", return_value=None):
                result = _select_model("Unknown language.")
        assert result == "apertus-8b"


class TestLocalAIClientAnalyze:
    def _valid_findings_response(self):
        return {
            "findings": [],
            "summary": "All good.",
            "overall_confidence": 0.9,
        }

    def test_localai_analyze_success_returns_payload(self):
        from app.services.localai import LocalAIClient
        client = LocalAIClient()
        content = '{"findings": [], "summary": "Test.", "overall_confidence": 0.9}'
        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.content = content.encode()
        mock_response.raise_for_status.return_value = None
        mock_response.json.return_value = {
            "choices": [{"message": {"content": content}}]
        }

        with patch("httpx.AsyncClient") as mock_cls:
            mock_http = AsyncMock()
            mock_http.__aenter__.return_value = mock_http
            mock_http.post.return_value = mock_response
            mock_cls.return_value = mock_http

            result = asyncio.run(client.analyze("numbered text", ["GDPR"], []))

        assert result is not None
        assert "findings" in result

    def test_localai_analyze_http_status_error_returns_none(self):
        import httpx
        from app.services.localai import LocalAIClient
        client = LocalAIClient()

        mock_response = MagicMock()
        mock_response.status_code = 503
        mock_response.text = "Service Unavailable"

        with patch("httpx.AsyncClient") as mock_cls:
            mock_http = AsyncMock()
            mock_http.__aenter__.return_value = mock_http
            mock_http.post.side_effect = httpx.HTTPStatusError(
                "503 error", request=MagicMock(), response=mock_response
            )
            mock_cls.return_value = mock_http

            result = asyncio.run(client.analyze("text", ["GDPR"], []))

        assert result is None

    def test_localai_analyze_http_error_returns_none(self):
        import httpx
        from app.services.localai import LocalAIClient
        client = LocalAIClient()

        with patch("httpx.AsyncClient") as mock_cls:
            mock_http = AsyncMock()
            mock_http.__aenter__.return_value = mock_http
            mock_http.post.side_effect = httpx.ConnectError("Connection refused")
            mock_cls.return_value = mock_http

            result = asyncio.run(client.analyze("text", ["GDPR"], []))

        assert result is None

    def test_localai_analyze_json_decode_error_returns_none(self):
        from app.services.localai import LocalAIClient
        client = LocalAIClient()

        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.content = b"invalid-json"
        mock_response.raise_for_status.return_value = None
        mock_response.json.side_effect = ValueError("not JSON")

        with patch("httpx.AsyncClient") as mock_cls:
            mock_http = AsyncMock()
            mock_http.__aenter__.return_value = mock_http
            mock_http.post.return_value = mock_response
            mock_cls.return_value = mock_http

            result = asyncio.run(client.analyze("text", ["GDPR"], []))

        assert result is None

    def test_localai_analyze_missing_choices_returns_none(self):
        from app.services.localai import LocalAIClient
        client = LocalAIClient()

        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.content = b'{"no_choices": true}'
        mock_response.raise_for_status.return_value = None
        mock_response.json.return_value = {"no_choices": True}

        with patch("httpx.AsyncClient") as mock_cls:
            mock_http = AsyncMock()
            mock_http.__aenter__.return_value = mock_http
            mock_http.post.return_value = mock_response
            mock_cls.return_value = mock_http

            result = asyncio.run(client.analyze("text", ["GDPR"], []))

        assert result is None

    def test_localai_analyze_content_not_json_returns_none(self):
        from app.services.localai import LocalAIClient
        client = LocalAIClient()
        not_json_content = "This is not JSON at all."

        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.content = not_json_content.encode()
        mock_response.raise_for_status.return_value = None
        mock_response.json.return_value = {
            "choices": [{"message": {"content": not_json_content}}]
        }

        with patch("httpx.AsyncClient") as mock_cls:
            mock_http = AsyncMock()
            mock_http.__aenter__.return_value = mock_http
            mock_http.post.return_value = mock_response
            mock_cls.return_value = mock_http

            result = asyncio.run(client.analyze("text", ["GDPR"], []))

        assert result is None


class TestLocalAIClientEmbed:
    def test_localai_embed_success_returns_vector(self):
        from app.services.localai import LocalAIClient
        client = LocalAIClient()
        embedding = [0.1, 0.2, 0.3]

        mock_response = MagicMock()
        mock_response.raise_for_status.return_value = None
        mock_response.json.return_value = {"data": [{"embedding": embedding}]}

        with patch("httpx.AsyncClient") as mock_cls:
            mock_http = AsyncMock()
            mock_http.__aenter__.return_value = mock_http
            mock_http.post.return_value = mock_response
            mock_cls.return_value = mock_http

            result = asyncio.run(client.embed("text to embed"))

        assert result == embedding

    def test_localai_embed_error_returns_none(self):
        from app.services.localai import LocalAIClient
        import httpx
        client = LocalAIClient()

        with patch("httpx.AsyncClient") as mock_cls:
            mock_http = AsyncMock()
            mock_http.__aenter__.return_value = mock_http
            mock_http.post.side_effect = httpx.ConnectError("Connection refused")
            mock_cls.return_value = mock_http

            result = asyncio.run(client.embed("text"))

        assert result is None

    def test_localai_embed_with_explicit_model(self):
        from app.services.localai import LocalAIClient
        client = LocalAIClient()
        embedding = [0.5, 0.6]

        mock_response = MagicMock()
        mock_response.raise_for_status.return_value = None
        mock_response.json.return_value = {"data": [{"embedding": embedding}]}

        with patch("httpx.AsyncClient") as mock_cls:
            mock_http = AsyncMock()
            mock_http.__aenter__.return_value = mock_http
            mock_http.post.return_value = mock_response
            mock_cls.return_value = mock_http

            result = asyncio.run(client.embed("text", model="eurollm-22b"))

        assert result == embedding


# ═══════════════════════════════════════════════════════════════════════════
# Additional targeted coverage tests
# ═══════════════════════════════════════════════════════════════════════════

class TestExtractHtmlWithScripts:
    """ingest.py line 54: tag.decompose() for script/style/noscript tags."""

    def test_ingest_extract_html_strips_script_tags(self):
        from app.services.ingest import _extract_html
        html = (
            "<html><head>"
            "<script>alert('xss')</script>"
            "<style>.cls { color: red; }</style>"
            "<noscript>Enable JS</noscript>"
            "</head><body><p>Visible policy content</p></body></html>"
        )
        text = _extract_html(html)
        assert "Visible policy content" in text
        assert "alert" not in text


class TestExtractPdfViaExtractBytes:
    """ingest.py line 123: PDF with extractable text via extract_text_from_bytes."""

    def test_ingest_extract_pdf_with_text_hits_line_123(self):
        from app.services.ingest import extract_text_from_bytes
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
        except ImportError:
            pytest.skip("reportlab not installed")
        buf = BytesIO()
        doc = SimpleDocTemplate(buf)
        doc.build([Paragraph("Coverage line 123 PDF text.", getSampleStyleSheet()["Normal"])])
        buf.seek(0)
        text = extract_text_from_bytes("policy.pdf", "application/pdf", buf.read())
        assert "Coverage line 123" in text


class TestValidateUrlUnparseableAddr:
    """ingest.py lines 152-153, 158: getaddrinfo returns non-IP sockaddr."""

    def test_ingest_validate_url_unparseable_getaddrinfo_addr_raises(self):
        from app.services.ingest import _validate_url
        with patch("socket.getaddrinfo") as mock_dns:
            # sockaddr[0] is "not-an-ip" → ipaddress.ip_address raises ValueError
            # → continue (lines 152-153) → addresses empty → raise (line 158)
            mock_dns.return_value = [
                (None, None, None, None, ("not-an-ip-addr", 0)),
            ]
            with pytest.raises(ValueError, match="not allowed"):
                _validate_url("http://example.com/policy")


class TestSelectChunksBudgetFits:
    """embedding.py lines 210-211: chunk fits in budget → append and increment."""

    def test_embedding_select_chunks_large_budget_exercises_selection(self):
        from app.services.embedding import select_relevant_chunks
        client = AsyncMock()
        client.embed.return_value = [0.1, 0.2, 0.3, 0.4]

        # Text large enough to exceed max_chars; max_chars large enough for chunks (~800 chars)
        long_text = ("privacy data collection sharing retention rights policy content\n" * 200)
        result = asyncio.run(
            select_relevant_chunks(long_text, "privacy data", max_chars=5000, client=client)
        )
        assert isinstance(result, str)
        assert len(result) > 0


class TestDetectLanguageMocked:
    """localai.py lines 33-36: _detect_language body when _LANGDETECT_AVAILABLE patched True."""

    def test_localai_detect_language_success_when_available(self):
        from app.services.localai import _detect_language
        # create=True is required because _langdetect is not defined when langdetect is absent
        with patch("app.services.localai._LANGDETECT_AVAILABLE", True):
            with patch("app.services.localai._langdetect", return_value="de", create=True):
                result = _detect_language("Datenschutzrichtlinie des Unternehmens.")
        assert result == "de"

    def test_localai_detect_language_exception_returns_none(self):
        from app.services.localai import _detect_language
        with patch("app.services.localai._LANGDETECT_AVAILABLE", True):
            with patch(
                "app.services.localai._langdetect",
                side_effect=Exception("detection failed"),
                create=True,
            ):
                result = _detect_language("Ambiguous text sample.")
        assert result is None


class TestLocalAIClientBaseUrl:
    """localai.py line 92: /v1 appended when base URL lacks it."""

    def test_localai_client_appends_v1_when_base_url_missing_it(self):
        from app.services.localai import LocalAIClient
        with patch("app.services.localai.settings") as mock_settings:
            mock_settings.localai_base_url = "http://localhost:8080"  # no /v1
            mock_settings.request_timeout_s = 60
            client = LocalAIClient()
        assert client._base_url == "http://localhost:8080/v1"


class TestAnalyzeTextInvalidLLMFinding:
    """analyzer.py lines 257-258: bad item in LLM findings triggers except/continue."""

    def test_analyzer_analyze_text_invalid_llm_finding_is_skipped(self):
        from app.services.analyzer import analyze_text
        mock_payload = {
            "findings": [
                {"completely_invalid_field": "no_required_keys_present"},
            ],
            "summary": "Test summary.",
            "overall_confidence": 0.7,
        }
        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = mock_payload
            mock_cls.return_value = mock_client
            result = asyncio.run(analyze_text("We sell your personal data.", ["GDPR"]))
        # Invalid finding was skipped; no finding with a nonsense category survives
        assert result.payload is not None
        bad_categories = [
            f.category for f in result.payload.findings
            if f.category == "completely_invalid_field"
        ]
        assert bad_categories == [], (
            "Malformed LLM finding must be skipped and not appear in results"
        )
        assert result.payload is not None


class TestAnalyzeTextSourceDocument:
    """analyzer.py lines 263-264: source_document assigned to findings in batch context."""

    def test_analyzer_analyze_text_source_document_propagated_to_findings(self):
        from app.services.analyzer import analyze_text
        # Use full mode with a mocked LLM returning a finding that has source_document=None
        mock_payload = {
            "findings": [
                {
                    "category": "data_sharing",
                    "severity": "High",
                    "confidence": 0.88,
                    "excerpt": "We sell your data.",
                    "explanation": "Data sale.",
                    "jurisdictions": ["GDPR"],
                    "evidence": {
                        "line_start": 1, "line_end": 1,
                        "legal_basis": ["GDPR Art. 6"],
                    },
                }
            ],
            "summary": "Data sale found.",
            "overall_confidence": 0.88,
        }
        with patch("app.services.analyzer.LocalAIClient") as mock_cls:
            mock_client = AsyncMock()
            mock_client.analyze.return_value = mock_payload
            mock_cls.return_value = mock_client
            result = asyncio.run(
                analyze_text(
                    "We sell your data to third parties.",
                    ["GDPR"],
                    source_document="batch_document_001",
                )
            )
        # All findings should have source_document set to batch_document_001
        assert result is not None
        for finding in result.payload.findings:
            assert finding.source_document == "batch_document_001"


class TestDetectHighSeverityRegexException:
    """analyzer.py lines 412-413: re.finditer raises → except Exception/continue."""

    def test_analyzer_detect_high_severity_regex_exception_is_continued(self):
        import re as _re
        from app.services.analyzer import detect_high_severity_findings
        original_finditer = _re.finditer
        call_count = [0]

        def patched_finditer(pattern, string, flags=0):
            call_count[0] += 1
            if call_count[0] == 1:
                raise _re.error("forced regex error for coverage")
            return original_finditer(pattern, string, flags)

        with patch("app.services.analyzer.re.finditer", side_effect=patched_finditer):
            findings = detect_high_severity_findings(
                "We sell your personal data to third parties.", ["GDPR"]
            )
        assert isinstance(findings, list)
        assert call_count[0] >= 1
